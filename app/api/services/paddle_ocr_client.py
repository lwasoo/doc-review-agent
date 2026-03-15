import asyncio
from pathlib import Path
from typing import Any, Dict, List
import sys
import types
import inspect
import os

import fitz
import numpy as np

from common.logger import get_logger
from config.config import settings

logging = get_logger(__name__)


class PaddleOCRClient:
    """
    Local OCR client based on PaddleOCR.
    Keeps interface similar to MinerUClient:
    - extract(file_path) -> {"content": ..., "meta": ...}
    - to_paragraphs(payload) -> list[paragraph]
    """

    def __init__(self) -> None:
        self._engine = None

    def _get_engine(self):
        if self._engine is not None:
            return self._engine
        _ensure_langchain_docstore_shim()
        os.environ.setdefault("PADDLE_PDX_DISABLE_MODEL_SOURCE_CHECK", "True")
        try:
            from paddleocr import PaddleOCR  # type: ignore
        except Exception as e:
            raise RuntimeError(
                "PaddleOCR import failed. Ensure compatible deps are installed. "
                f"Original error: {e}"
            ) from e

        init_params = set(inspect.signature(PaddleOCR.__init__).parameters.keys())
        kwargs: Dict[str, Any] = {}
        if "lang" in init_params:
            kwargs["lang"] = settings.paddleocr_lang
        if "show_log" in init_params:
            kwargs["show_log"] = False
        if "use_angle_cls" in init_params:
            kwargs["use_angle_cls"] = bool(settings.paddleocr_use_angle_cls)
        elif "use_textline_orientation" in init_params:
            kwargs["use_textline_orientation"] = bool(settings.paddleocr_use_angle_cls)

        init_errors: List[str] = []
        for attempt in (kwargs, {"lang": settings.paddleocr_lang}, {}):
            try:
                self._engine = PaddleOCR(**attempt)
                break
            except Exception as e:
                init_errors.append(f"{attempt}: {e}")
                self._engine = None
        if self._engine is None:
            raise RuntimeError(
                "PaddleOCR init failed. Attempts: " + " | ".join(init_errors)
            )
        return self._engine

    async def extract(self, file_path: Path) -> Dict[str, Any]:
        if not file_path.exists():
            raise FileNotFoundError(str(file_path))
        return await asyncio.to_thread(self._extract_sync, file_path)

    def _extract_sync(self, file_path: Path) -> Dict[str, Any]:
        ext = file_path.suffix.lower()
        paragraphs: List[Dict[str, Any]] = []

        if ext == ".pdf":
            engine = self._get_engine()
            doc = fitz.open(str(file_path))
            try:
                for i in range(doc.page_count):
                    page = doc.load_page(i)
                    pix = page.get_pixmap(
                        dpi=max(96, int(settings.paddleocr_pdf_dpi)),
                        alpha=False,
                    )
                    img = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.h, pix.w, pix.n)
                    if pix.n >= 3:
                        img = img[:, :, :3]

                    result = _run_ocr(engine, img, bool(settings.paddleocr_use_angle_cls))
                    lines = _extract_ocr_lines(result)
                    for line in lines:
                        text = str(line.get("text", "")).strip()
                        if not text:
                            continue
                        paragraphs.append(
                            {
                                "content": text,
                                "page_num": i + 1,
                                "bbox": line.get("bbox"),
                                "page_height": float(pix.h),
                                "canvas_size": [int(pix.w), int(pix.h)],
                            }
                        )
            finally:
                doc.close()
        elif ext in {".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff", ".webp"}:
            engine = self._get_engine()
            result = _run_ocr(engine, str(file_path), bool(settings.paddleocr_use_angle_cls))
            lines = _extract_ocr_lines(result)
            for line in lines:
                text = str(line.get("text", "")).strip()
                if not text:
                    continue
                paragraphs.append(
                    {
                        "content": text,
                        "page_num": 1,
                        "bbox": line.get("bbox"),
                        "page_height": None,
                        "canvas_size": None,
                    }
                )
        elif ext == ".docx":
            paragraphs = _extract_docx_paragraphs(file_path)
        elif ext == ".doc":
            raise RuntimeError(
                "PaddleOCR provider does not directly support .doc. "
                "Please convert to .docx or .pdf first."
            )
        else:
            raise RuntimeError(
                f"PaddleOCR provider currently supports PDF/images/docx, got: {file_path.suffix}"
            )

        return {
            "content": {"paragraphs": paragraphs},
            "meta": {"engine": "paddleocr", "page_count": max([p["page_num"] for p in paragraphs], default=0)},
        }

    @staticmethod
    def to_paragraphs(payload: Dict[str, Any]) -> List[Dict[str, Any]]:
        content = payload.get("content") if isinstance(payload, dict) else None
        if isinstance(content, dict):
            paras = content.get("paragraphs")
            if isinstance(paras, list):
                return paras
        if isinstance(payload, dict) and isinstance(payload.get("paragraphs"), list):
            return payload["paragraphs"]
        return []


def _extract_ocr_lines(result: Any) -> List[Dict[str, Any]]:
    """
    Normalize PaddleOCR output into:
    [{"text": str, "score": float|None, "bbox": [x0,y0,x1,y1]}]
    """
    lines: List[Dict[str, Any]] = []
    for item in _iter_paddle_line_items(result):
        try:
            points = item[0]
            info = item[1]
            text = info[0] if isinstance(info, (list, tuple)) and len(info) > 0 else ""
            score = info[1] if isinstance(info, (list, tuple)) and len(info) > 1 else None
            bbox = _points_to_bbox(points)
            lines.append({"text": text, "score": score, "bbox": bbox})
        except Exception:
            continue
    return lines


def _iter_paddle_line_items(obj: Any):
    if isinstance(obj, (list, tuple)):
        if _is_line_item(obj):
            yield obj
            return
        for sub in obj:
            yield from _iter_paddle_line_items(sub)


def _is_line_item(obj: Any) -> bool:
    if not isinstance(obj, (list, tuple)) or len(obj) < 2:
        return False
    points = obj[0]
    info = obj[1]
    if not isinstance(points, (list, tuple)) or not points:
        return False
    if not isinstance(info, (list, tuple)) or not info:
        return False
    return isinstance(info[0], str)


def _points_to_bbox(points: Any) -> List[float] | None:
    try:
        xs = [float(p[0]) for p in points]
        ys = [float(p[1]) for p in points]
        return [min(xs), min(ys), max(xs), max(ys)]
    except Exception:
        return None


def _ensure_langchain_docstore_shim() -> None:
    """
    PaddleX in some PaddleOCR versions imports:
    `from langchain.docstore.document import Document`
    `from langchain.text_splitter import RecursiveCharacterTextSplitter`
    but LangChain 1.x removed that path.
    Provide a minimal compatibility shim at runtime.
    """
    if (
        "langchain.docstore.document" in sys.modules
        and "langchain.text_splitter" in sys.modules
    ):
        return
    try:
        from langchain_core.documents import Document  # type: ignore
    except Exception:
        return

    doc_mod = types.ModuleType("langchain.docstore.document")
    setattr(doc_mod, "Document", Document)

    docstore_mod = types.ModuleType("langchain.docstore")
    setattr(docstore_mod, "document", doc_mod)

    splitter_mod = types.ModuleType("langchain.text_splitter")
    try:
        from langchain_text_splitters import RecursiveCharacterTextSplitter  # type: ignore
    except Exception:
        class RecursiveCharacterTextSplitter:  # type: ignore
            def __init__(self, chunk_size: int = 300, chunk_overlap: int = 20, separators=None):
                self.chunk_size = max(1, int(chunk_size))
                self.chunk_overlap = max(0, int(chunk_overlap))
                self.separators = separators or ["\n\n", "\n", " ", ""]

            def split_text(self, text: str) -> list[str]:
                text = text or ""
                if not text:
                    return []
                out: list[str] = []
                step = max(1, self.chunk_size - self.chunk_overlap)
                i = 0
                while i < len(text):
                    out.append(text[i : i + self.chunk_size])
                    i += step
                return out
    setattr(splitter_mod, "RecursiveCharacterTextSplitter", RecursiveCharacterTextSplitter)

    # If top-level langchain module exists, attach .docstore for attribute access.
    langchain_mod = sys.modules.get("langchain")
    if langchain_mod is not None:
        if not hasattr(langchain_mod, "docstore"):
            setattr(langchain_mod, "docstore", docstore_mod)
        if not hasattr(langchain_mod, "text_splitter"):
            setattr(langchain_mod, "text_splitter", splitter_mod)

    sys.modules["langchain.docstore"] = docstore_mod
    sys.modules["langchain.docstore.document"] = doc_mod
    sys.modules["langchain.text_splitter"] = splitter_mod


def _extract_docx_paragraphs(file_path: Path) -> List[Dict[str, Any]]:
    try:
        from docx import Document  # type: ignore
    except Exception as e:
        raise RuntimeError(
            "DOCX parsing dependency missing. Please run: pip install python-docx"
        ) from e

    doc = Document(str(file_path))
    paragraphs: List[Dict[str, Any]] = []

    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if text:
            paragraphs.append(
                {
                    "content": text,
                    "page_num": 1,
                    "bbox": None,
                    "page_height": None,
                    "canvas_size": None,
                }
            )

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = "\n".join([(cp.text or "").strip() for cp in cell.paragraphs]).strip()
                if text:
                    paragraphs.append(
                        {
                            "content": text,
                            "page_num": 1,
                            "bbox": None,
                            "page_height": None,
                            "canvas_size": None,
                        }
                    )

    return paragraphs


def _run_ocr(engine: Any, inp: Any, use_angle_cls: bool) -> Any:
    try:
        return engine.ocr(inp, cls=use_angle_cls)
    except TypeError:
        # PaddleOCR 3.x may not accept cls in this method.
        return engine.ocr(inp)
    except NotImplementedError as e:
        msg = str(e)
        if "ConvertPirAttribute2RuntimeAttribute" in msg:
            raise RuntimeError(
                "PaddleOCR runtime is incompatible in current environment. "
                "Please use OCR_PROVIDER=mineru, or run API in a Python 3.10/3.11 "
                "environment with a PaddleOCR/PaddlePaddle-compatible version."
            ) from e
        raise
