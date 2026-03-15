from __future__ import annotations

from datetime import datetime
from pathlib import Path
import re
from typing import Iterable, List

from common.models import Issue, IssueStatusEnum
from config.config import settings


def export_review_docx(source_path: Path, issues: List[Issue], accepted_only: bool = True) -> Path:
    from docx import Document  # type: ignore
    from docx.shared import RGBColor  # type: ignore

    export_dir = Path(settings.export_docs_dir)
    export_dir.mkdir(parents=True, exist_ok=True)

    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    out_name = f"{source_path.stem}_审阅版_{ts}.docx"
    out_path = export_dir / out_name

    filtered = list(_filter_issues(issues, accepted_only=accepted_only))

    if source_path.suffix.lower() == ".docx" and source_path.exists():
        doc = Document(str(source_path))
        _append_inline_suggestions(doc, filtered)
    else:
        doc = Document()
        doc.add_heading("合同审阅报告（DOCX导出）", level=1)
        doc.add_paragraph(f"源文件：{source_path.name}")
        doc.add_paragraph("说明：原文件非 DOCX，以下为问题与建议清单。")

    _append_review_summary(doc, source_path.name, filtered, accepted_only)
    doc.save(str(out_path))
    return out_path


def _filter_issues(issues: List[Issue], accepted_only: bool) -> Iterable[Issue]:
    if not accepted_only:
        return [i for i in issues if i.status != IssueStatusEnum.dismissed]
    return [i for i in issues if i.status == IssueStatusEnum.accepted]


def _append_inline_suggestions(doc, issues: List[Issue]) -> None:
    from docx.shared import RGBColor  # type: ignore

    paras = list(doc.paragraphs)
    for idx, issue in enumerate(issues, 1):
        needle = _normalize(issue.text or "")
        if not needle:
            continue
        hit = None
        for p in paras:
            pt = _normalize(p.text or "")
            if not pt:
                continue
            if needle in pt or pt in needle:
                hit = p
                break
            if _similar(needle, pt) > 0.82:
                hit = p
                break
        if hit is None:
            continue
        note = doc.add_paragraph()
        run = note.add_run(f"【审阅建议#{idx} | {issue.type}】{issue.suggested_fix}")
        run.font.color.rgb = RGBColor(192, 0, 0)


def _append_review_summary(doc, src_name: str, issues: List[Issue], accepted_only: bool) -> None:
    doc.add_page_break()
    doc.add_heading("审阅建议清单", level=1)
    doc.add_paragraph(f"源文件：{src_name}")
    doc.add_paragraph(f"导出范围：{'仅采纳建议' if accepted_only else '未驳回建议'}")
    doc.add_paragraph(f"问题数量：{len(issues)}")

    if not issues:
        doc.add_paragraph("无可导出建议。")
        return

    table = doc.add_table(rows=1, cols=5)
    hdr = table.rows[0].cells
    hdr[0].text = "序号"
    hdr[1].text = "风险类型"
    hdr[2].text = "等级"
    hdr[3].text = "原文片段"
    hdr[4].text = "建议条款"

    for i, issue in enumerate(issues, 1):
        row = table.add_row().cells
        row[0].text = str(i)
        row[1].text = issue.type or ""
        row[2].text = str(issue.risk_level or "")
        row[3].text = (issue.text or "")[:800]
        row[4].text = issue.suggested_fix or ""


def _normalize(s: str) -> str:
    return re.sub(r"\s+", "", (s or "").strip().lower())


def _similar(a: str, b: str) -> float:
    if not a or not b:
        return 0.0
    sa = set(a)
    sb = set(b)
    inter = len(sa & sb)
    den = max(len(sa), len(sb), 1)
    return inter / den
